# book_style.py
# Shared publication framework for The Guardians of Elaria chronicle.
# Encodes the Session 8 house style: US Letter, parchment-gold read-aloud
# boxes, lavender DM boxes, goldenrod/sienna headings, star dividers,
# gray italic captions, and a centered star page footer.
# Build scripts import build_doc() and feed it a list of content blocks.

import io
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# House palette (from Session_8_Gearhaven_v2)
INK = "222222"          # body text
GOLD_FILL = "FBF6EA"    # read-aloud box fill
GOLD_EDGE = "B8860B"    # read-aloud left rule + H1 color
PURPLE_FILL = "F4F0FA"  # DM box fill
PURPLE_EDGE = "5B2A86"  # DM box left rule (Ursa purple)
STAT_FILL = "FCF6F6"    # statblock fill
STAT_EDGE = "8B2020"    # statblock left rule (deep crimson)
H2_COLOR = "A0522D"     # sienna
H3_COLOR = "1F4D78"     # deep blue
CAPTION_GRAY = "888888"
ACCENTS = {"lilly": "1F6FB8", "stabby": "A32B2B", "ursa": "5B2A86",
           "ghostbloom": "1F7A78", "gold": GOLD_EDGE}

BODY_FONT = "Georgia"
BODY_SIZE = Pt(10.5)

def _set_font(run, size=None, bold=None, italic=None, color=None, font=BODY_FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)
    if size: run.font.size = size
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color or INK)

def _shade(paragraph, fill, edge):
    """Left rule + shading, matching the S8 box construction."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '12'); left.set(qn('w:color'), edge)
    pbdr.append(left); ppr.append(pbdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    ppr.append(shd)
    paragraph.paragraph_format.left_indent = Twips(340)
    paragraph.paragraph_format.right_indent = Twips(340)
    paragraph.paragraph_format.space_before = Twips(30)
    paragraph.paragraph_format.space_after = Twips(70)

def _rich(par, text, base_color=INK, base_size=BODY_SIZE, base_bold=False, base_italic=False):
    """Minimal inline markup: **bold**, *italic*, and {color:name}...{/} accents."""
    import re
    pos = 0
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|\{color:\w+\}.*?\{/\})', text)
    for tok in tokens:
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); _set_font(r, base_size, True, base_italic, base_color)
        elif tok.startswith('{color:'):
            import re as _re
            m = _re.match(r'\{color:(\w+)\}(.*)\{/\}', tok, _re.S)
            r = par.add_run(m.group(2)); _set_font(r, base_size, True, base_italic, ACCENTS.get(m.group(1), base_color))
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            r = par.add_run(tok[1:-1]); _set_font(r, base_size, base_bold, True, base_color)
        else:
            r = par.add_run(tok); _set_font(r, base_size, base_bold, base_italic, base_color)

def _footer_stars(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("✦ "); _set_font(r1, Pt(9), color=CAPTION_GRAY)
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    rf = p.add_run(); _set_font(rf, Pt(9), color=CAPTION_GRAY)
    rf._element.append(fld1)
    rm = p.add_run(); _set_font(rm, Pt(9), color=CAPTION_GRAY)
    rm._element.append(instr)
    re_ = p.add_run(); _set_font(re_, Pt(9), color=CAPTION_GRAY)
    re_._element.append(fld2)
    r2 = p.add_run(" ✦"); _set_font(r2, Pt(9), color=CAPTION_GRAY)

def _image_png_bytes(path, max_w_px=1200, crop=None):
    """Return compressed JPEG bytes for any raster (webp converted, large images
    downscaled). JPEG keeps the illustrated docx/PDF a sane size; the art is
    opaque painterly work, so no alpha is lost.
    crop: optional "W:H" string (e.g. "4:3"). Square art cropped to a landscape
    ratio renders as a full-width plate instead of a tall centered square."""
    from PIL import Image
    im = Image.open(path).convert("RGB")
    if crop:
        wr, hr = (float(x) for x in crop.split(":"))
        target = hr / wr
        cur = im.height / im.width
        if cur > target + 0.01:
            # too tall: trim height, biased slightly above center (skies crop
            # better than foregrounds)
            nh = int(im.width * target)
            top = int((im.height - nh) * 0.42)
            im = im.crop((0, top, im.width, top + nh))
        elif cur < target - 0.01:
            nw = int(im.height / target)
            left = (im.width - nw) // 2
            im = im.crop((left, 0, left + nw, im.height))
    if im.width > max_w_px:
        im = im.resize((max_w_px, int(im.height * max_w_px / im.width)), Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=82, optimize=True)
    return buf.getvalue(), im.width, im.height

def _sb_line(doc, label, text, size=Pt(9)):
    p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r = p.add_run(label + " "); _set_font(r, size, True, color=STAT_EDGE)
    _rich(p, text, base_size=size)

def _sb_rule(doc):
    p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    ppr = p._p.get_or_add_pPr(); pbdr = ppr.find(qn('w:pBdr'))
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4'); bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), STAT_EDGE)
    pbdr.append(bot)

def _mod(score):
    m = (int(score) - 10) // 2
    return f"+{m}" if m >= 0 else str(m)

def _float_right(paragraph, data, w_in):
    """Anchor a picture floated to the right with square text-wrap, so the
    statblock body wraps around it instead of stranding whitespace."""
    from docx.oxml import parse_xml
    run = paragraph.add_run()
    pic = run.add_picture(io.BytesIO(data), width=Inches(w_in))
    cx, cy = pic._inline.extent.cx, pic._inline.extent.cy
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    graphic = inline.find(qn('a:graphic'))
    docpr = inline.find(qn('wp:docPr'))
    anchor = parse_xml(
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'behindDoc="0" distT="91440" distB="100584" distL="201600" distR="201600" simplePos="0" locked="0" '
        'layoutInCell="1" allowOverlap="0" relativeHeight="2">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:align>right</wp:align></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapSquare wrapText="bothSides"/>'
        f'<wp:docPr id="{docpr.get("id")}" name="{docpr.get("name")}"/>'
        '<wp:cNvGraphicFramePr/>'
        '</wp:anchor>')
    anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)


def _render_statblock(doc, sb):
    """A 5e-style statblock in the crimson house box, with the creature
    portrait floated to the right and the text wrapping around it. sb keys:
    name, type, ac, hp, speed, abilities {STR..CHA}, optional cr, saves,
    skills, resistances, vulnerabilities, immunities, condition_immunities,
    senses, languages, traits/actions/reactions/legendary [(name,text)], img, img_w."""
    # name header, and float the portrait from this first paragraph so the
    # whole block wraps around it
    _sb_start = len(doc.paragraphs)
    p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
    p.paragraph_format.keep_with_next = True; p.paragraph_format.space_after = Pt(0)
    if sb.get("img"):
        try:
            data, pw, ph = _image_png_bytes(sb["img"])
            w = sb.get("img_w", 2.35)
            if w * (ph / pw) > 2.6:
                w = 2.6 / (ph / pw)
            _float_right(p, data, w)
        except Exception:
            pass
    r = p.add_run(sb["name"]); _set_font(r, Pt(12.5), True, color=STAT_EDGE)
    # type line
    p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(sb["type"]); _set_font(r, Pt(8.5), italic=True, color=INK)
    _sb_rule(doc)
    if sb.get("ac"): _sb_line(doc, "Armor Class", sb["ac"])
    if sb.get("hp"): _sb_line(doc, "Hit Points", sb["hp"])
    if sb.get("speed"): _sb_line(doc, "Speed", sb["speed"])
    _sb_rule(doc)
    # ability scores as a single wrapping line (plays nicely beside the float)
    ab = sb.get("abilities")
    if ab:
        p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
        p.paragraph_format.space_after = Pt(1)
        for i, name in enumerate(["STR", "DEX", "CON", "INT", "WIS", "CHA"]):
            sc = ab.get(name, 10)
            rr = p.add_run(("   " if i else "") + name + " "); _set_font(rr, Pt(9), True, color=STAT_EDGE)
            rr = p.add_run(f"{sc} ({_mod(sc)})"); _set_font(rr, Pt(9), color=INK)
        _sb_rule(doc)
    for key, label in [("saves", "Saving Throws"), ("skills", "Skills"),
                       ("resistances", "Damage Resistances"), ("vulnerabilities", "Damage Vulnerabilities"),
                       ("immunities", "Damage Immunities"), ("condition_immunities", "Condition Immunities"),
                       ("senses", "Senses"), ("languages", "Languages"), ("cr", "Challenge")]:
        if sb.get(key):
            _sb_line(doc, label, sb[key])
    for section, items in [("Traits", sb.get("traits")), ("Actions", sb.get("actions")),
                           ("Reactions", sb.get("reactions")), ("Legendary Actions", sb.get("legendary"))]:
        if not items:
            continue
        if section != "Traits":
            _sb_rule(doc)
            p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
            r = p.add_run(section); _set_font(r, Pt(10), True, color=STAT_EDGE)
        for nm, txt in items:
            p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
            p.paragraph_format.space_after = Pt(1)
            if nm:
                r = p.add_run(nm + ". "); _set_font(r, Pt(9), True, italic=True, color=INK)
            _rich(p, txt, base_size=Pt(9))
    # chain the block's head together so a statblock can never START in the
    # last sliver of a page and sink its floated portrait off the bottom edge
    for _kp in doc.paragraphs[_sb_start:_sb_start + 8]:
        _kp.paragraph_format.keep_with_next = True
        _kp.paragraph_format.keep_together = True
    # a trailing spacer so the next block clears the floated image if it overhangs
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _page_parchment(doc, hexcolor="F9F2E2"):
    """Book-wide parchment page color (exports to PDF via Word)."""
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), hexcolor)
    doc.element.insert(0, bg)
    settings = doc.settings.element
    if settings.find(qn('w:displayBackgroundShape')) is None:
        el = OxmlElement('w:displayBackgroundShape')
        settings.append(el)


def build_doc(blocks, out_path):
    doc = Document()
    _page_parchment(doc)
    # US Letter, S8 margins
    sec = doc.sections[0]
    sec.page_width = Twips(12240); sec.page_height = Twips(15840)
    sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    sec.top_margin = Twips(1080); sec.bottom_margin = Twips(1300)
    _footer_stars(sec)

    st_normal = doc.styles['Normal']
    st_normal.font.name = BODY_FONT
    st_normal.font.size = BODY_SIZE
    st_normal.font.color.rgb = RGBColor.from_string(INK)

    for _bi, blk in enumerate(blocks):
        kind = blk[0]
        _next_kind = blocks[_bi + 1][0] if _bi + 1 < len(blocks) else None

        if kind == "titlepage":
            # (titlepage, kicker, title, subtitle, note)
            _, kicker, title, subtitle, note = blk
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            r = p.add_run("✦  ✦  ✦"); _set_font(r, Pt(16), color=GOLD_EDGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(kicker); _set_font(r, Pt(15), True, color=INK)
            r.font.name = BODY_FONT
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title); _set_font(r, Pt(30), True, color=GOLD_EDGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle); _set_font(r, Pt(17), True, color=H2_COLOR)
            if note:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(16)
                r = p.add_run(note); _set_font(r, Pt(9.5), italic=True, color=CAPTION_GRAY)

        elif kind == "h1":
            p = doc.add_paragraph()
            # Content flows continuously; only hard-break before the Appendix
            # (or when a block explicitly asks with hardbreak). Parts are set
            # off by generous space and a keep-with-next rule, not a page break.
            hard = blk[1].strip().lower().startswith("appendix")
            if len(blk) > 2 and blk[2].get("hardbreak"):
                hard = True
            if hard:
                p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(blk[1]); _set_font(r, Pt(16), True, color=GOLD_EDGE)
            # Chapter-header rule: a thin goldenrod line under every part title,
            # the touch that reads as a 5e-style sourcebook heading.
            _ppr = p._p.get_or_add_pPr()
            _pbdr = OxmlElement('w:pBdr')
            _bot = OxmlElement('w:bottom')
            _bot.set(qn('w:val'), 'single'); _bot.set(qn('w:sz'), '6')
            _bot.set(qn('w:space'), '4'); _bot.set(qn('w:color'), GOLD_EDGE)
            _pbdr.append(_bot); _ppr.append(_pbdr)

        elif kind == "h2":
            # (h2, title[, {"hardbreak": True}]) - hardbreak sets Word's
            # page-break-before property, which (unlike an explicit break
            # paragraph) is suppressed when the heading already opens a page,
            # so repagination can never mint a blank page.
            p = doc.add_paragraph()
            if len(blk) > 2 and isinstance(blk[2], dict) and blk[2].get("hardbreak"):
                p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(blk[1]); _set_font(r, Pt(13), True, color=H2_COLOR)

        elif kind == "gold":
            p = doc.add_paragraph(); _shade(p, GOLD_FILL, GOLD_EDGE)
            _rich(p, blk[1])

        elif kind == "dm":
            # (dm, text[, {"size": pt}]) - optional size for lines meant to be
            # read at the table rather than skimmed (character-sheet stats).
            dsz = blk[2].get("size", 9.5) if len(blk) > 2 and isinstance(blk[2], dict) else 9.5
            p = doc.add_paragraph(); _shade(p, PURPLE_FILL, PURPLE_EDGE)
            _rich(p, "▶ " + blk[1], base_size=Pt(dsz))

        elif kind == "slots":
            # (slots, accent_key, [(label, count), ...]) - a character-sheet
            # style resource tracker: big filled pips in the hero's accent
            # color, one row per slot level. Reads at arm's length.
            _, who, rows = blk
            acc = ACCENTS.get(who, INK)
            for i, (label, n) in enumerate(rows):
                p = doc.add_paragraph(); _shade(p, GOLD_FILL, acc)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2 if i < len(rows) - 1 else 8)
                r = p.add_run(label + "   "); _set_font(r, Pt(11.5), True, color=acc)
                r = p.add_run("● " * int(n)); _set_font(r, Pt(14), True, color=acc)

        elif kind == "spellref":
            # (spellref, accent_key, {"note": str, "levels": [(label, slots, [names]), ...]})
            # A compact "Spells at a Glance" play aid for a caster's sheet: a
            # check-off slot tracker (open boxes) beside the spell NAMES known
            # at each level. Complements the pips (glance the count) and the
            # full spellbook (rules); this is the mid-fight tracker plus index.
            _, who, spec = blk
            acc = ACCENTS.get(who, INK)
            levels = spec["levels"]
            p = doc.add_paragraph(); _shade(p, GOLD_FILL, acc)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run("Spells at a Glance"); _set_font(r, Pt(12.5), True, color=acc)
            if spec.get("note"):
                r = p.add_run("   " + spec["note"]); _set_font(r, Pt(9), italic=True, color=INK)
            for i, (label, slots, sp) in enumerate(levels):
                p = doc.add_paragraph(); _shade(p, GOLD_FILL, acc)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2 if i < len(levels) - 1 else 8)
                p.paragraph_format.keep_with_next = i < len(levels) - 1
                r = p.add_run(label + "  "); _set_font(r, Pt(11), True, color=acc)
                if slots and int(slots) > 0:
                    r = p.add_run("☐ " * int(slots)); _set_font(r, Pt(12.5), True, color=acc)
                else:
                    r = p.add_run("at will"); _set_font(r, Pt(9.5), italic=True, color=INK)
                r = p.add_run("   " + ", ".join(sp)); _set_font(r, Pt(10), color=INK)

        elif kind == "game":
            # (game, title, [lines]) - a SET-PIECE RULES CARD for playable
            # table moments (dice games, song seals, storm sequences). A full
            # double-ruled teal frame, big title, roomy text: the one box in
            # the book that says "put the book down and play this".
            _, gtitle, glines = blk[0], blk[1], blk[2]
            G_FILL, G_EDGE = "EAF4F3", "1F7A78"
            def _card(par, first=False, last=False):
                ppr = par._p.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                for side, on in [("top", first), ("bottom", last), ("left", True), ("right", True)]:
                    if not on:
                        continue
                    el = OxmlElement('w:' + side)
                    el.set(qn('w:val'), 'double'); el.set(qn('w:sz'), '12')
                    el.set(qn('w:space'), '10' if side in ("left", "right") else '4')
                    el.set(qn('w:color'), G_EDGE)
                    pbdr.append(el)
                ppr.append(pbdr)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), G_FILL)
                ppr.append(shd)
                par.paragraph_format.left_indent = Twips(400)
                par.paragraph_format.right_indent = Twips(400)
            p = doc.add_paragraph(); _card(p, first=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run("✦  " + gtitle + "  ✦"); _set_font(r, Pt(13.5), True, color=G_EDGE)
            for i, line in enumerate(glines):
                p = doc.add_paragraph(); _card(p, last=(i == len(glines) - 1))
                p.paragraph_format.space_before = Pt(0)
                # zero spacing INSIDE the card so the fill reads as one
                # continuous panel; the frame closes after the last line.
                p.paragraph_format.space_after = Pt(14 if i == len(glines) - 1 else 0)
                # a card NEVER splits across a page turn: every line keeps
                # its own lines together and chains to the next line.
                p.paragraph_format.keep_together = True
                if i < len(glines) - 1:
                    p.paragraph_format.keep_with_next = True
                _rich(p, line, base_size=Pt(10.5))

        elif kind == "stat":
            # (stat, title, [lines])
            p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(blk[1]); _set_font(r, Pt(11), True, color=STAT_EDGE)
            for line in blk[2]:
                p = doc.add_paragraph(); _shade(p, STAT_FILL, STAT_EDGE)
                _rich(p, line, base_size=Pt(9))

        elif kind == "statblock":
            _render_statblock(doc, blk[1])

        elif kind == "ua_stat":
            # (ua_stat, {name, type_line, top: [(label, text)], abilities:
            # [(ab, score)], meta: [(label, text)], sections: [(header,
            # [(entry, text)])]}) - the official summon-spell block anatomy
            # (DM directive 2026-07-13): full-width banded name/section
            # headers, labeled rows with hairline rules, a six-column
            # ability table, and "(Partner Only)" annotated entries.
            spec = blk[1]

            def _ua_band(text, size=Pt(10.5)):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), GOLD_EDGE)
                pPr.append(shd)
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(text); _set_font(r, size, True, color="FFFFFF")
                return p

            def _ua_row(rich_text, keep=True):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                bdr = OxmlElement('w:pBdr')
                el = OxmlElement('w:bottom')
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
                el.set(qn('w:space'), '2'); el.set(qn('w:color'), 'E0D8C4')
                bdr.append(el); pPr.append(bdr)
                if keep:
                    p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
                _rich(p, rich_text, base_size=Pt(9.5))
                return p

            _ua_band(spec["name"])
            _ua_row("*" + spec.get("type_line", "") + "*")
            for label, text in spec.get("top", []):
                _ua_row("**" + label + ":** " + text)
            abilities = spec.get("abilities", [])
            if abilities:
                atbl = doc.add_table(rows=2, cols=len(abilities))
                atbl.alignment = 1
                for j, (ab, score) in enumerate(abilities):
                    hc = atbl.rows[0].cells[j]
                    hp_ = hc.paragraphs[0]
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), GOLD_EDGE)
                    hc._tc.get_or_add_tcPr().append(shd)
                    r = hp_.add_run(ab); _set_font(r, Pt(9), True, color="FFFFFF")
                    vc = atbl.rows[1].cells[j]
                    vp = vc.paragraphs[0]
                    vp.paragraph_format.space_after = Pt(0)
                    r = vp.add_run(score); _set_font(r, Pt(9.5))
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(0)
                rr = sp.add_run(""); _set_font(rr, Pt(2))
            for label, text in spec.get("meta", []):
                _ua_row("**" + label + ":** " + text)
            for header, entries in spec.get("sections", []):
                # The band keeps with its first entry; entries flow freely so a
                # long section can split across a page turn like official print.
                _ua_band(header, size=Pt(10))
                for ename, etext in entries:
                    _ua_row("**" + ename + ".** " + etext, keep=False)

        elif kind == "body":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _rich(p, blk[1])

        elif kind == "bridge":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
            _rich(p, blk[1], base_italic=True)

        elif kind == "hero":
            # (hero, who, line) suggested PC voice, accent colored
            _, who, line = blk
            p = doc.add_paragraph(); _shade(p, GOLD_FILL, GOLD_EDGE)
            r = p.add_run(who + ": "); _set_font(r, BODY_SIZE, True, color=ACCENTS.get(who.split(' ')[0].lower(), INK))
            _rich(p, "“" + line + "”", base_italic=True)

        elif kind == "divider":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(10)
            r = p.add_run("✦ ✦ ✦"); _set_font(r, Pt(12), color=GOLD_EDGE)

        elif kind == "img":
            # (img, path, caption, width_inches[, opts])
            # opts: {"crop": "4:3", "hmax": 5.2} - crop square art to a
            # landscape plate; hmax raises the height cap for showpieces.
            _, path, caption, w = blk[0], blk[1], blk[2], blk[3]
            opts = blk[4] if len(blk) > 4 and isinstance(blk[4], dict) else {}
            data, pw, ph = _image_png_bytes(path, crop=opts.get("crop"))
            # Cap rendered HEIGHT so tall/portrait art cannot eat a whole page
            # and strand whitespace. Landscape art keeps its requested width.
            aspect = ph / pw
            max_h = opts.get("hmax", 4.2)
            if w * aspect > max_h:
                w = max_h / aspect
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = bool(caption)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(io.BytesIO(data), width=Inches(w))
            if caption:
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(7)
                r = cp.add_run(caption); _set_font(r, Pt(9), italic=True, color=CAPTION_GRAY)

        elif kind == "imgrow":
            # (imgrow, [(path, label), (path, label)], width_each) side-by-side
            _, pairs, w = blk
            tbl = doc.add_table(rows=2, cols=len(pairs))
            tbl.autofit = False
            tbl.alignment = 1
            # zero the side cell margins: Word's default ~0.08in per side
            # pushed wide rows a hair past the right text margin.
            tpr = tbl._tbl.tblPr
            mar = OxmlElement('w:tblCellMar')
            for side in ('left', 'right'):
                el = OxmlElement('w:' + side)
                el.set(qn('w:w'), '0'); el.set(qn('w:type'), 'dxa')
                mar.append(el)
            tpr.append(mar)
            for ci, (path, label) in enumerate(pairs):
                data, pw, ph = _image_png_bytes(path)
                ww = w
                if ww * (ph / pw) > 3.6:
                    ww = 3.6 / (ph / pw)
                c0 = tbl.cell(0, ci); c1 = tbl.cell(1, ci)
                c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # keep the image row on the same page as its caption row
                c0.paragraphs[0].paragraph_format.keep_with_next = True
                c0.paragraphs[0].add_run().add_picture(io.BytesIO(data), width=Inches(ww))
                c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = c1.paragraphs[0].add_run(label)
                _set_font(r, Pt(9), italic=True, color=CAPTION_GRAY)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif kind == "railrow":
            # (railrow, rail_path, rail_w_in, [inner blocks]) - a text column
            # beside a stacked image rail, rendered as ONE table row. This is
            # the engine-safe replacement for anchored floats: it flows with
            # the page, splits across pages, and can never overlap a margin
            # or shred a wrap. Inner blocks support h2 / gold / dm / body.
            # Rail images carry their own baked-in captions.
            _, rpath, rw, inner = blk
            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            tbl.alignment = 1
            tpr = tbl._tbl.tblPr
            mar = OxmlElement('w:tblCellMar')
            for side in ('left', 'right'):
                el = OxmlElement('w:' + side)
                el.set(qn('w:w'), '0'); el.set(qn('w:type'), 'dxa')
                mar.append(el)
            tpr.append(mar)
            text_w = 6.5 - rw - 0.22
            ct, cimg = tbl.cell(0, 0), tbl.cell(0, 1)
            ct.width = Inches(text_w); cimg.width = Inches(rw + 0.22)
            first = True
            for ib in inner:
                ikind = ib[0]
                if first:
                    p = ct.paragraphs[0]; first = False
                else:
                    p = ct.add_paragraph()
                p.paragraph_format.right_indent = Pt(10)
                if ikind == "h2":
                    p.paragraph_format.space_before = Pt(9)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.keep_with_next = True
                    r = p.add_run(ib[1]); _set_font(r, Pt(13), True, color=H2_COLOR)
                elif ikind == "gold":
                    _shade(p, GOLD_FILL, GOLD_EDGE); _rich(p, ib[1])
                elif ikind == "dm":
                    _shade(p, PURPLE_FILL, PURPLE_EDGE)
                    _rich(p, "▶ " + ib[1], base_size=Pt(9.5))
                else:
                    _rich(p, ib[1])
            data, pw, ph = _image_png_bytes(rpath)
            cimg.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cimg.paragraphs[0].add_run().add_picture(io.BytesIO(data), width=Inches(rw))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif kind == "melody":
            # (melody, text) a sung verse or carol, set like sheet-music epigraph
            for i, line in enumerate(blk[1].split("|")):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10 if i == 0 else 0)
                p.paragraph_format.space_after = Pt(2)
                txt = line.strip()
                if i == 0:
                    txt = "♪  " + txt
                if i == len(blk[1].split("|")) - 1:
                    txt = txt + "  ♪"
                r = p.add_run(txt.replace("\u266a", "♪"))
                _set_font(r, Pt(12), bold=False, italic=True, color=GOLD_EDGE)
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

        elif kind == "imgfloat":
            # (imgfloat, path, width_inches[, "left"|"right"]) - image floated
            # into the following text with square wrap; no caption. The text
            # that comes AFTER this block flows around the image.
            _, path, w = blk[0], blk[1], blk[2]
            side = blk[3] if len(blk) > 3 and isinstance(blk[3], str) else "right"
            opts = next((b for b in blk[3:] if isinstance(b, dict)), {})
            data, pw, ph = _image_png_bytes(path, crop=opts.get("crop"))
            fmax = opts.get("hmax", 3.4)
            if w * (ph / pw) > fmax:
                w = fmax / (ph / pw)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
            _float_right(p, data, w)
            if side == "left":
                # flip the anchor alignment to left
                anch = p._p.findall('.//' + qn('wp:anchor'))
                if anch:
                    al = anch[0].find(qn('wp:positionH') + '/' + qn('wp:align'))
                    if al is None:
                        ph_el = anch[0].find(qn('wp:positionH'))
                        al = ph_el.find(qn('wp:align'))
                    if al is not None:
                        al.text = 'left'

        elif kind == "lore":
            # (lore, title, text) - a "Lore of Elaria" fragment: world history
            # in the chronicler's voice, dropped onto art-heavy pages. Deep
            # blue (the one unused house color), thin double rules above and
            # below, small caps kicker. Reads as a margin note from an older
            # book. Fragments come from lore/world_history.md only.
            _, ltitle, ltext = blk
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.left_indent = Twips(500); p.paragraph_format.right_indent = Twips(500)
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'double'); top.set(qn('w:sz'), '6')
            top.set(qn('w:space'), '4'); top.set(qn('w:color'), H3_COLOR)
            pbdr.append(top); ppr.append(pbdr)
            r = p.add_run("✦ LORE OF ELARIA ✦  ")
            _set_font(r, Pt(8), True, color=H3_COLOR)
            r = p.add_run(ltitle)
            _set_font(r, Pt(9.5), True, color=H3_COLOR)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.left_indent = Twips(500); p.paragraph_format.right_indent = Twips(500)
            p.paragraph_format.keep_together = True
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'double'); bot.set(qn('w:sz'), '6')
            bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), H3_COLOR)
            pbdr.append(bot); ppr.append(pbdr)
            _rich(p, ltext, base_size=Pt(9.5), base_italic=True)

        elif kind == "appendix_title":
            # (appendix_title, kicker, title, subtitle) - the S7 appendix
            # opener: centered gold kicker, big crimson title, sienna
            # subtitle, on a fresh page.
            _, kick, atitle, asub = blk
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(kick); _set_font(r, Pt(10.5), True, color=GOLD_EDGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.keep_with_next = True
            r = p.add_run(atitle); _set_font(r, Pt(16), True, color=STAT_EDGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(asub); _set_font(r, Pt(10.5), color=H2_COLOR)

        elif kind == "fight_header":
            # (fight_header, title, subline[, {"compact": True}]) - "ENEMIES,
            # Fight 1: ..." over a gray location/difficulty/roster line, both
            # centered (S7 style). compact=True lowers the space_before and
            # drops the subline's keep_with_next so a fight can begin low on a
            # page under a preceding short card row instead of reserving a whole
            # block and forcing a page break (2026-07-09 whitespace pass).
            _, ftitle, fsub = blk[0], blk[1], blk[2]
            fopts = blk[3] if len(blk) > 3 and isinstance(blk[3], dict) else {}
            compact = fopts.get("compact", False)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10 if compact else 16)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(ftitle); _set_font(r, Pt(13), True, color=STAT_EDGE)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = not compact
            r = p.add_run(fsub); _set_font(r, Pt(9), color=CAPTION_GRAY)

        elif kind == "enemy_cards":
            # (enemy_cards, [card, ...][, {"pack": True}]) - 1 to 3 S7-style
            # bestiary cards side by side (DM directive 2026-07-07: multi-type
            # encounters sit side by side; bosses go solo full-width). card:
            # {name, sub, img (optional), stats: [lines], traits: [(n,t)],
            # actions: [(n,t)], reactions: [(n,t)]}. Lines support **bold** /
            # *italic*. pack=True compacts a minor-mob reference row (smaller
            # thumbnails) so several rows fit one page (2026-07-09 whitespace).
            _, cards = blk[0], blk[1]
            ecopts = blk[2] if len(blk) > 2 and isinstance(blk[2], dict) else {}
            packmode = ecopts.get("pack", False)
            ncols = max(1, min(3, len(cards)))
            tbl = doc.add_table(rows=1, cols=ncols)
            tbl.autofit = False
            tbl.alignment = 1
            tpr = tbl._tbl.tblPr
            mar = OxmlElement('w:tblCellMar')
            for side, wd in (('left', '110'), ('right', '110'), ('top', '60'), ('bottom', '80')):
                el = OxmlElement('w:' + side)
                el.set(qn('w:w'), wd); el.set(qn('w:type'), 'dxa')
                mar.append(el)
            tpr.append(mar)
            borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement('w:' + side)
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '8')
                el.set(qn('w:space'), '0'); el.set(qn('w:color'), STAT_EDGE)
                borders.append(el)
            tpr.append(borders)
            col_in = 6.3 / ncols
            # A SOLO card (one enemy, usually a boss or companion) renders its
            # portrait as a big centered PLATE below the statblock, so the card
            # fills the page and the art gets the size the DM keeps asking for,
            # instead of a small side float that strands a tail (2026-07-09
            # whitespace pass). A card may force plate=False to keep the float.
            # Multi-card rows keep the centered-above look.
            solo = (ncols == 1)

            def _bodylen(c):
                n = len(c.get("stats", []))
                for _k in ("traits", "actions", "reactions", "legendary"):
                    _it = c.get(_k)
                    if _it:
                        n += 1 + len(_it)
                return n

            solo_plate = bool(solo and cards[0].get("img") and cards[0].get("plate", True))
            # A solo card with a long statblock PLUS a full plate can exceed one
            # page; let that (and only that) card split at a paragraph boundary so
            # it fills two pages rather than jumping wholesale and stranding a tail.
            allow_split = bool(solo_plate and _bodylen(cards[0]) >= 22)
            # Every other card row never splits across a page turn.
            if not allow_split:
                trpr = tbl.rows[0]._tr.get_or_add_trPr()
                trpr.append(OxmlElement('w:cantSplit'))
            for ci, card in enumerate(cards[:ncols]):
                cell = tbl.cell(0, ci)
                cell.width = Inches(col_in)
                # banner
                bp = cell.paragraphs[0]
                _cshd = OxmlElement('w:shd')
                _cshd.set(qn('w:val'), 'clear'); _cshd.set(qn('w:color'), 'auto')
                _cshd.set(qn('w:fill'), card.get("banner", STAT_EDGE))
                bp._p.get_or_add_pPr().append(_cshd)
                bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(1)
                r = bp.add_run(card["name"].upper()); _set_font(r, Pt(11), True, color="FFFFFF")
                # type line
                p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(card.get("sub", "")); _set_font(r, Pt(8), italic=True, color=STAT_EDGE)
                solo_float = None
                plate = None
                if card.get("img"):
                    try:
                        data, pw, ph = _image_png_bytes(card["img"], crop=card.get("crop"))
                        if solo and solo_plate:
                            # big plate rendered BELOW the stats (see loop end)
                            plate = (data, pw, ph)
                        elif solo:
                            # side float, cap raised from 2.8 to 3.4 (2026-07-09)
                            solo_float = (data, min(card.get("img_w", 2.5), 3.4))
                        else:
                            cap = (col_in * 0.5) if packmode else (col_in - 0.15)
                            iw = min(card.get("img_w", cap), cap)
                            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(3)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run().add_picture(io.BytesIO(data), width=Inches(iw))
                    except Exception:
                        pass
                first_wrap = True
                sz = Pt(8.5) if solo else Pt(7.5)
                for line in card.get("stats", []):
                    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1)
                    if solo_float and first_wrap:
                        _float_right(p, solo_float[0], solo_float[1]); first_wrap = False
                    _rich(p, line, base_size=sz)
                for sect, key in (("Traits", "traits"), ("Actions", "actions"),
                                  ("Reactions", "reactions"), ("Legendary Actions", "legendary")):
                    items = card.get(key)
                    if not items:
                        continue
                    p = cell.add_paragraph(); p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(1); p.paragraph_format.keep_with_next = True
                    if solo_float and first_wrap:
                        _float_right(p, solo_float[0], solo_float[1]); first_wrap = False
                    r = p.add_run(sect); _set_font(r, Pt(9 if solo else 8), True, color=STAT_EDGE)
                    for nm, txt in items:
                        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1)
                        if nm:
                            r = p.add_run(nm + ". "); _set_font(r, sz, True, italic=True, color=INK)
                        _rich(p, txt, base_size=sz)
                # SOLO plate: a big centered portrait below the whole statblock.
                # Landscape art becomes a wide half-page plate; square/portrait
                # art a generous centered figure. Height is capped so a near
                # square plate cannot re-strand a page below it.
                if plate:
                    pdata, ppw, pph = plate
                    aspect = pph / ppw  # rendered height / width
                    landscape = ppw >= pph * 1.15
                    # Moderate plate below the statblock: landscape art becomes a
                    # wide half-page plate, square/portrait a generous centered
                    # figure (floored so a small legacy img_w still fills). Height
                    # is capped so a card stays under one page and never orphans
                    # the following creature's intro line (2026-07-09).
                    if landscape:
                        pw_in = min(col_in - 0.25, 5.7)
                        if pw_in * aspect > 3.5:
                            pw_in = 3.5 / aspect
                    else:
                        pw_in = min(max(card.get("plate_w", card.get("img_w", 3.6)), 3.0), 4.1)
                        if pw_in * aspect > 4.1:
                            pw_in = 4.1 / aspect
                    pp = cell.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pp.paragraph_format.space_before = Pt(5); pp.paragraph_format.space_after = Pt(1)
                    pp.add_run().add_picture(io.BytesIO(pdata), width=Inches(pw_in))
            # No trailing spacer paragraph: it strands a blank page when a card
            # group ends exactly at a page boundary before a hardbreak header.
            # Inter-group separation comes from the following fight_header /
            # h1 spacing; the table's own bottom cell margin handles the rest.

        elif kind == "tactics":
            # (tactics, text) - the S7 "Tactics:" line under a fight's cards.
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
            r = p.add_run("Tactics: "); _set_font(r, Pt(9), True, color=STAT_EDGE)
            _rich(p, blk[1], base_size=Pt(9), base_italic=True)

        elif kind == "reward_card":
            # (reward_card, title, sub, [(item, desc), ...][, img]) - the gold
            # sibling of an enemy card: a full-width in-session item rewards
            # block (banner in goldenrod, item names bold).
            _, rtitle, rsub, ritems = blk[0], blk[1], blk[2], blk[3]
            rimg = blk[4] if len(blk) > 4 else None
            tbl = doc.add_table(rows=1, cols=1)
            tbl.autofit = False
            tbl.alignment = 1
            tpr = tbl._tbl.tblPr
            mar = OxmlElement('w:tblCellMar')
            for side, wd in (('left', '110'), ('right', '110'), ('top', '60'), ('bottom', '80')):
                el = OxmlElement('w:' + side)
                el.set(qn('w:w'), wd); el.set(qn('w:type'), 'dxa')
                mar.append(el)
            tpr.append(mar)
            borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement('w:' + side)
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '8')
                el.set(qn('w:space'), '0'); el.set(qn('w:color'), GOLD_EDGE)
                borders.append(el)
            tpr.append(borders)
            trpr = tbl.rows[0]._tr.get_or_add_trPr()
            trpr.append(OxmlElement('w:cantSplit'))
            cell = tbl.cell(0, 0)
            cell.width = Inches(6.3)
            bp = cell.paragraphs[0]
            _cshd = OxmlElement('w:shd')
            _cshd.set(qn('w:val'), 'clear'); _cshd.set(qn('w:color'), 'auto')
            _cshd.set(qn('w:fill'), GOLD_EDGE)
            bp._p.get_or_add_pPr().append(_cshd)
            bp.paragraph_format.space_before = Pt(0); bp.paragraph_format.space_after = Pt(1)
            r = bp.add_run(rtitle.upper()); _set_font(r, Pt(11), True, color="FFFFFF")
            p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(rsub); _set_font(r, Pt(8), italic=True, color=H2_COLOR)
            if rimg:
                try:
                    data, pw, ph = _image_png_bytes(rimg)
                    iw = 4.6
                    if iw * (ph / pw) > 3.2:
                        iw = 3.2 / (ph / pw)
                    p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(3)
                    p.add_run().add_picture(io.BytesIO(data), width=Inches(iw))
                except Exception:
                    pass
            for nm, txt in ritems:
                p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(2)
                if nm:
                    r = p.add_run(nm + ". "); _set_font(r, Pt(9.5), True, color=STAT_EDGE)
                _rich(p, txt, base_size=Pt(9.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif kind == "pagebreak":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

        else:
            raise ValueError(f"unknown block kind {kind}")

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
